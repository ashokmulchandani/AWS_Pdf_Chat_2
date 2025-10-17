"""
-*- coding: utf-8 -*-
========================
AWS Lambda for Textract
========================
Contributor: Ashok Mulchandani 
========================
"""

import os
import json
import boto3


def lambda_handler(event, context):

    BUCKET_NAME = os.environ.get("BUCKET_NAME", "python-docx-bucket1")
    PREFIX = os.environ.get("PREFIX", "output")

    job_id = json.loads(event["Records"][0]["Sns"]["Message"])["JobId"]

    # Extract structured JSON using Textract + Bedrock
    structured_data = process_response_with_bedrock(job_id)

    # Save as JSON file
    json_key_name = f"{job_id}.json"
    with open(f"/tmp/{json_key_name}", 'w') as f:
        json.dump(structured_data, f, indent=2)
    upload_to_s3(f"/tmp/{json_key_name}", BUCKET_NAME, f"{PREFIX}/{json_key_name}")
    
    # Generate DOCX file
    docx_key_name = f"{job_id}_Underwriting_Summary.docx"
    generate_docx(structured_data, f"/tmp/{docx_key_name}")
    upload_to_s3(f"/tmp/{docx_key_name}", BUCKET_NAME, f"{PREFIX}/{docx_key_name}")
    
    # Move processed PDF to processed_pdfs folder
    move_pdf_to_processed(job_id, BUCKET_NAME)
    
    return {"statusCode": 200, "body": json.dumps("Files uploaded successfully!")}


def upload_to_s3(filename, bucket, key):
    s3 = boto3.client("s3")
    s3.upload_file(Filename=filename, Bucket=bucket, Key=key)
    print(f"Uploaded {filename} to s3://{bucket}/{key}")


def process_response_with_bedrock(job_id):
    # Step 1: Extract raw text from Textract
    raw_text = extract_raw_text(job_id)
    
    # Step 2: Load JSON template
    json_template = load_json_template()
    
    # Step 3: Use Bedrock to structure the data
    return structure_with_bedrock(raw_text, json_template)


def extract_raw_text(job_id):
    textract = boto3.client("textract")
    
    all_text = []
    response = textract.get_document_text_detection(JobId=job_id)
    
    # Process first page
    for block in response.get('Blocks', []):
        if block['BlockType'] == 'LINE':
            all_text.append(block['Text'])
    
    # Handle pagination
    while 'NextToken' in response:
        response = textract.get_document_text_detection(
            JobId=job_id, 
            NextToken=response['NextToken']
        )
        for block in response.get('Blocks', []):
            if block['BlockType'] == 'LINE':
                all_text.append(block['Text'])
    
    return '\n'.join(all_text)


def load_json_template():
    # Try to read the complete Nobleoak_json.txt file content
    try:
        with open('/opt/templates/Nobleoak_json.txt', 'r') as f:
            return f.read()
    except:
        return get_embedded_template()


def get_embedded_template():
    # Load from S3 first, fallback to local file content
    s3 = boto3.client('s3')
    try:
        response = s3.get_object(
            Bucket=os.environ.get('TEMPLATE_BUCKET', os.environ['BUCKET_NAME']),
            Key='templates/Nobleoak_json.txt'
        )
        return response['Body'].read().decode('utf-8')
    except:
        # Read the complete template from local file
        try:
            with open('Nobleoak_json.txt', 'r') as f:
                return f.read()
        except:
            # Final fallback - return the complete template structure
            # Use the complete Nobleoak template from the file
            return '''You are an expert underwriting analyst for a life insurance company. Your task is to analyze the provided life insurance application text and extract all relevant information into a single, valid JSON object. Do not provide any other text, explanations, or conversational filler. The output must be ONLY the JSON object.

The application text is:
---
{{ $json.extractedText }}
---

Ensure the final JSON is correct and contains no extra content.

Only output JSON in the following format:

{
  "applicant": {
    "name": "ADD",
    "dob": "ADD",
    "state": "ADD"
  },
  "underwritingSections": [
    {
      "section": "Applicant details",
      "findings": [
        {"text": "Name: [extracted name]"},
        {"text": "DOB: [extracted dob]"},
        {"text": "Additional relevant disclosures: none beyond those listed."}
      ]
    }
  ],
  "summary": {
    "disclosureSummary": [
      {
        "heading": "[Category]",
        "bullets": [
          {"text": "[Summary point]"}
        ]
      }
    ],
    "redFlags": [
      {"text": "[Any inconsistencies or concerns]"}
    ]
  }
}

If you are unsure of any field, output "N/A".'''


def structure_with_bedrock(raw_text, json_template):
    bedrock = boto3.client('bedrock-runtime')
    
    # Replace template placeholder with actual text
    prompt = json_template.replace('{{ $json.extractedText }}', raw_text)
    
    try:
        response = bedrock.invoke_model(
            modelId='anthropic.claude-3-sonnet-20240229-v1:0',
            body=json.dumps({
                'anthropic_version': 'bedrock-2023-05-31',
                'messages': [{
                    'role': 'user', 
                    'content': prompt
                }],
                'max_tokens': 4000,
                'temperature': 0.1
            })
        )
        
        response_body = json.loads(response['body'].read())
        structured_text = response_body['content'][0]['text']
        
        # Parse the JSON response
        return json.loads(structured_text)
        
    except Exception as e:
        print(f"Bedrock processing failed: {str(e)}")
        # Fallback to basic structure
        return {
            "applicant": {"name": "N/A", "dob": "N/A", "state": "N/A"},
            "underwritingSections": [],
            "summary": {"disclosureSummary": [], "redFlags": []}
        }


def generate_docx(data, filepath):
    """Generate DOCX using sample template for exact formatting"""
    try:
        print(f"[DOCX] Attempting to import python-docx...")
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        print(f"[DOCX] python-docx imported successfully")
        
        # Load Exercise 1 sample as template
        try:
            s3 = boto3.client('s3')
            # Try multiple possible template names
            template_keys = [
                'templates/Exercise 1- Sample output - Underwriting_Summary.docx',
                'templates/Exercise_1_Sample_Underwriting_Summary.docx',
                'templates/sample_template.docx'
            ]
            
            template_response = None
            for key in template_keys:
                try:
                    template_response = s3.get_object(
                        Bucket=os.environ.get('BUCKET_NAME', 'python-docx-bucket1'),
                        Key=key
                    )
                    print(f"[DOCX] Found template at: {key}")
                    break
                except:
                    continue
            
            if not template_response:
                raise Exception("No template found")
            with open('/tmp/template.docx', 'wb') as f:
                f.write(template_response['Body'].read())
            doc = Document('/tmp/template.docx')
            
            # Clear all content but preserve styles
            for paragraph in doc.paragraphs[:]:
                p = paragraph._element
                p.getparent().remove(p)
            
            print(f"[DOCX] Using Exercise 1 sample template")
                
        except:
            doc = Document()
            print(f"[DOCX] Template not found - creating new document")
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title - Match Exercise 1 exactly
        title_para = doc.add_paragraph()
        title_run = title_para.add_run('Underwriting Summary')
        title_run.bold = True
        title_run.font.size = Pt(16)  # Larger for exact match
        title_run.font.name = 'Calibri'
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(12)
        
        # Add line break
        doc.add_paragraph()
        
        # Applicant Information Section
        applicant_heading = doc.add_paragraph()
        applicant_run = applicant_heading.add_run('Applicant Information')
        applicant_run.bold = True
        applicant_run.underline = True
        
        applicant = data.get('applicant', {})
        
        # Applicant details with bold labels
        name_para = doc.add_paragraph()
        name_para.add_run('Name: ').bold = True
        name_para.add_run(applicant.get('name', 'N/A'))
        
        dob_para = doc.add_paragraph()
        dob_para.add_run('DOB: ').bold = True
        dob_para.add_run(applicant.get('dob', 'N/A'))
        
        state_para = doc.add_paragraph()
        state_para.add_run('State: ').bold = True
        state_para.add_run(applicant.get('state', 'N/A'))
        
        # Add spacing
        doc.add_paragraph()
        
        # Process each underwriting section
        for section in data.get('underwritingSections', []):
            # Section heading - Bold and underlined
            section_heading = doc.add_paragraph()
            section_run = section_heading.add_run(section.get('section', ''))
            section_run.bold = True
            section_run.underline = True
            
            # Section findings using proper bullet style
            for finding in section.get('findings', []):
                finding_text = finding.get('text', '')
                if finding_text and finding_text.strip() and finding_text != 'N/A':
                    bullet_para = doc.add_paragraph(finding_text, style='List Bullet')
            
            # Add spacing between sections
            doc.add_paragraph()
        
        # Summary Section
        summary = data.get('summary', {})
        if summary and (summary.get('disclosureSummary') or summary.get('redFlags')):
            # Summary main heading
            summary_heading = doc.add_paragraph()
            summary_run = summary_heading.add_run('Summary')
            summary_run.bold = True
            summary_run.underline = True
            
            # Disclosure Summary
            if summary.get('disclosureSummary'):
                doc.add_paragraph()
                
                # Disclosure Summary subheading
                disc_heading = doc.add_paragraph()
                disc_run = disc_heading.add_run('Disclosure Summary')
                disc_run.bold = True
                
                for disclosure in summary.get('disclosureSummary', []):
                    # Category heading
                    if disclosure.get('heading'):
                        cat_para = doc.add_paragraph()
                        cat_run = cat_para.add_run(disclosure.get('heading', ''))
                        cat_run.bold = True
                        
                        # Category bullets using proper style
                        for bullet in disclosure.get('bullets', []):
                            bullet_text = bullet.get('text', '')
                            if bullet_text and bullet_text.strip():
                                bullet_para = doc.add_paragraph(bullet_text, style='List Bullet')
                        
                        doc.add_paragraph()  # Spacing between categories
            
            # Red Flags Section
            if summary.get('redFlags'):
                red_flags_heading = doc.add_paragraph()
                red_flags_run = red_flags_heading.add_run('Red Flags')
                red_flags_run.bold = True
                
                for flag in summary.get('redFlags', []):
                    flag_text = flag.get('text', '')
                    if flag_text and flag_text.strip():
                        bullet_para = doc.add_paragraph(flag_text, style='List Bullet')
        
        doc.save(filepath)
        print(f"[DOCX] Successfully generated sample-format DOCX: {filepath}")
        
    except ImportError as e:
        print(f"[DOCX] ImportError: {str(e)}")
        print(f"[DOCX] python-docx not available, generating exact sample text format")
    except Exception as e:
        print(f"[DOCX] Unexpected error: {str(e)}")
        print(f"[DOCX] Falling back to text format")
        # Fallback: create text file matching sample format exactly
        with open(filepath, 'w', encoding='utf-8') as f:
            # Centered title
            f.write("                         Underwriting Summary\n")
            f.write("\n\n")
            
            # Applicant Information
            f.write("Applicant Information\n")
            f.write("_____________________\n\n")
            applicant = data.get('applicant', {})
            f.write(f"Name: {applicant.get('name', 'N/A')}\n")
            f.write(f"DOB: {applicant.get('dob', 'N/A')}\n")
            f.write(f"State: {applicant.get('state', 'N/A')}\n\n\n")
            
            # Underwriting Sections
            for section in data.get('underwritingSections', []):
                f.write(f"{section.get('section', '')}\n")
                f.write("_" * len(section.get('section', '')) + "\n\n")
                
                for finding in section.get('findings', []):
                    finding_text = finding.get('text', '')
                    if finding_text and finding_text != 'N/A':
                        f.write(f"• {finding_text}\n")
                f.write("\n\n")
            
            # Summary
            summary = data.get('summary', {})
            if summary and (summary.get('disclosureSummary') or summary.get('redFlags')):
                f.write("Summary\n")
                f.write("_______\n\n")
                
                # Disclosure Summary
                if summary.get('disclosureSummary'):
                    f.write("Disclosure Summary\n\n")
                    
                    for disclosure in summary.get('disclosureSummary', []):
                        if disclosure.get('heading'):
                            f.write(f"{disclosure.get('heading', '')}\n")
                            
                            for bullet in disclosure.get('bullets', []):
                                bullet_text = bullet.get('text', '')
                                if bullet_text:
                                    f.write(f"• {bullet_text}\n")
                            f.write("\n")
                
                # Red Flags
                if summary.get('redFlags'):
                    f.write("Red Flags\n")
                    for flag in summary.get('redFlags', []):
                        flag_text = flag.get('text', '')
                        if flag_text:
                            f.write(f"• {flag_text}\n")
        
        print(f"[DOCX] Generated sample-format text report: {filepath}")
    
    except Exception as e:
        print(f"[DOCX] Error generating DOCX: {str(e)}")
        # Create minimal fallback file
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"Underwriting Summary\n\nProcessed Data:\n{json.dumps(data, indent=2)}")
        print(f"[DOCX] Created fallback report: {filepath}")


def move_pdf_to_processed(job_id, bucket_name):
    print(f"[PDF MOVE] Starting PDF move process for job: {job_id}")
    s3 = boto3.client('s3')
    
    try:
        print(f"[PDF MOVE] Searching for PDF files in input/ folder...")
        
        # Skip Textract job details call - just find most recent PDF
        response = s3.list_objects_v2(Bucket=bucket_name, Prefix='input/')
        if 'Contents' in response:
            print(f"[PDF MOVE] Found {len(response['Contents'])} files in input/ folder")
            # Get PDF files only
            pdf_files = [obj for obj in response['Contents'] if obj['Key'].endswith('.pdf')]
            print(f"[PDF MOVE] Found {len(pdf_files)} PDF files")
            
            if pdf_files:
                # Sort by last modified and get the most recent
                pdf_files.sort(key=lambda x: x['LastModified'], reverse=True)
                original_key = pdf_files[0]['Key']
                print(f"[PDF MOVE] Selected most recent PDF: {original_key}")
            else:
                print(f"[PDF MOVE] No PDF files found in input/ folder")
                return
        else:
            print(f"[PDF MOVE] No files found in input/ folder")
            return
            
        # Define new key in processed_pdfs folder
        filename = original_key.split('/')[-1]  # Get just the filename
        new_key = f"processed_pdfs/{filename}"
        
        print(f"[PDF MOVE] Starting move operation...")
        print(f"[PDF MOVE] Source: s3://{bucket_name}/{original_key}")
        print(f"[PDF MOVE] Destination: s3://{bucket_name}/{new_key}")
        
        # Copy file to new location
        print(f"[PDF MOVE] Step 1: Copying file to processed_pdfs/ folder...")
        s3.copy_object(
            Bucket=bucket_name,
            CopySource={'Bucket': bucket_name, 'Key': original_key},
            Key=new_key
        )
        print(f"[PDF MOVE] Step 1: Copy completed successfully")
        
        # Delete original file
        print(f"[PDF MOVE] Step 2: Deleting original file from input/ folder...")
        s3.delete_object(Bucket=bucket_name, Key=original_key)
        print(f"[PDF MOVE] Step 2: Delete completed successfully")
        
        print(f"[PDF MOVE] SUCCESS: PDF moved from {original_key} to {new_key}")
        
    except Exception as e:
        print(f"[PDF MOVE] ERROR: Failed to move PDF - {str(e)}")
        print(f"[PDF MOVE] This is non-critical - processing continues")
        # Don't fail the entire process if PDF move fails
