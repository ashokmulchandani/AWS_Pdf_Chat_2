"""
-*- coding: utf-8 -*-
========================
AWS Lambda for Textract
========================
Contributor: Ashok Mulchandani 
========================
"""

import os
import json
import boto3
from docx import Document
from docx.shared import Inches


def lambda_handler(event, context):

    BUCKET_NAME = os.environ.get("BUCKET_NAME", "python-docx-bucket1")
    PREFIX = os.environ.get("PREFIX", "output")

    job_id = json.loads(event["Records"][0]["Sns"]["Message"])["JobId"]

    # Extract structured JSON using Textract + Bedrock
    structured_data = process_response_with_bedrock(job_id)

    # Save as JSON file
    json_key_name = f"{job_id}.json"
    with open(f"/tmp/{json_key_name}", 'w') as f:
        json.dump(structured_data, f, indent=2)
    upload_to_s3(f"/tmp/{json_key_name}", BUCKET_NAME, f"{PREFIX}/{json_key_name}")
    
    # Generate DOCX file
    docx_key_name = f"{job_id}_Underwriting_Summary.docx"
    generate_docx(structured_data, f"/tmp/{docx_key_name}")
    upload_to_s3(f"/tmp/{docx_key_name}", BUCKET_NAME, f"{PREFIX}/{docx_key_name}")
    
    # Move processed PDF to processed_pdfs folder
    move_pdf_to_processed(job_id, BUCKET_NAME)
    
    return {"statusCode": 200, "body": json.dumps("Files uploaded successfully!")}


def upload_to_s3(filename, bucket, key):
    s3 = boto3.client("s3")
    s3.upload_file(Filename=filename, Bucket=bucket, Key=key)
    print(f"Uploaded {filename} to s3://{bucket}/{key}")


def process_response_with_bedrock(job_id):
    # Step 1: Extract raw text from Textract
    raw_text = extract_raw_text(job_id)
    
    # Step 2: Load JSON template
    json_template = load_json_template()
    
    # Step 3: Use Bedrock to structure the data
    return structure_with_bedrock(raw_text, json_template)


def extract_raw_text(job_id):
    textract = boto3.client("textract")
    
    all_text = []
    response = textract.get_document_text_detection(JobId=job_id)
    
    # Process first page
    for block in response.get('Blocks', []):
        if block['BlockType'] == 'LINE':
            all_text.append(block['Text'])
    
    # Handle pagination
    while 'NextToken' in response:
        response = textract.get_document_text_detection(
            JobId=job_id, 
            NextToken=response['NextToken']
        )
        for block in response.get('Blocks', []):
            if block['BlockType'] == 'LINE':
                all_text.append(block['Text'])
    
    return '\n'.join(all_text)


def load_json_template():
    # Try to read the complete Nobleoak_json.txt file content
    try:
        with open('/opt/templates/Nobleoak_json.txt', 'r') as f:
            return f.read()
    except:
        return get_embedded_template()


def get_embedded_template():
    # Load from S3 first, fallback to local file content
    s3 = boto3.client('s3')
    try:
        response = s3.get_object(
            Bucket=os.environ.get('TEMPLATE_BUCKET', os.environ['BUCKET_NAME']),
            Key='templates/Nobleoak_json.txt'
        )
        return response['Body'].read().decode('utf-8')
    except:
        # Read the complete template from local file
        try:
            with open('Nobleoak_json.txt', 'r') as f:
                return f.read()
        except:
            # Final fallback - return the complete template structure
            return '''You are an expert underwriting analyst for a life insurance company. Your task is to analyze the provided life insurance application text and extract all relevant information into a single, valid JSON object. Do not provide any other text, explanations, or conversational filler. The output must be ONLY the JSON object.

The application text is:
---
{{ $json.extractedText }}
---

Ensure the final JSON is correct and contains no extra content.

Only output JSON with ALL 17 sections: Applicant details, Applied for cover, Existing cover, Modified Terms, Claims history, Residency details, Occupation details, Income details, Travel, Recreation, Alcohol, Drug Use, Smoking, BMI, Medical History, Family History, GP Details.

If you are unsure of any field, output "N/A".'''


def structure_with_bedrock(raw_text, json_template):
    bedrock = boto3.client('bedrock-runtime')
    
    # Replace template placeholder with actual text
    prompt = json_template.replace('{{ $json.extractedText }}', raw_text)
    
    try:
        response = bedrock.invoke_model(
            modelId='anthropic.claude-3-sonnet-20240229-v1:0',
            body=json.dumps({
                'anthropic_version': 'bedrock-2023-05-31',
                'messages': [{
                    'role': 'user', 
                    'content': prompt
                }],
                'max_tokens': 4000,
                'temperature': 0.1
            })
        )
        
        response_body = json.loads(response['body'].read())
        structured_text = response_body['content'][0]['text']
        
        # Parse the JSON response
        return json.loads(structured_text)
        
    except Exception as e:
        print(f"Bedrock processing failed: {str(e)}")
        # Fallback to basic structure
        return {
            "applicant": {"name": "N/A", "dob": "N/A", "state": "N/A"},
            "underwritingSections": [],
            "summary": {"disclosureSummary": [], "redFlags": []}
        }


def generate_docx(data, filepath):
    doc = Document()
    
    # Title
    title = doc.add_heading('Underwriting Summary', 0)
    
    # Applicant Info
    doc.add_heading('Applicant Information', level=1)
    applicant = data.get('applicant', {})
    doc.add_paragraph(f"Name: {applicant.get('name', 'N/A')}")
    doc.add_paragraph(f"DOB: {applicant.get('dob', 'N/A')}")
    doc.add_paragraph(f"State: {applicant.get('state', 'N/A')}")
    
    # Underwriting Sections
    doc.add_heading('Underwriting Analysis', level=1)
    for section in data.get('underwritingSections', []):
        doc.add_heading(section.get('section', ''), level=2)
        for finding in section.get('findings', []):
            doc.add_paragraph(finding.get('text', ''), style='List Bullet')
    
    # Summary
    summary = data.get('summary', {})
    doc.add_heading('Summary', level=1)
    
    # Disclosure Summary
    doc.add_heading('Disclosure Summary', level=2)
    for disclosure in summary.get('disclosureSummary', []):
        doc.add_heading(disclosure.get('heading', ''), level=3)
        for bullet in disclosure.get('bullets', []):
            doc.add_paragraph(bullet.get('text', ''), style='List Bullet')
    
    # Red Flags
    if summary.get('redFlags'):
        doc.add_heading('Red Flags', level=2)
        for flag in summary.get('redFlags', []):
            doc.add_paragraph(flag.get('text', ''), style='List Bullet')
    
    doc.save(filepath)


def move_pdf_to_processed(job_id, bucket_name):
    s3 = boto3.client('s3')
    
    try:
        # Get the original PDF file path from Textract job
        textract = boto3.client('textract')
        job_response = textract.get_document_text_detection(JobId=job_id)
        
        # Extract original S3 key from job details
        document_location = job_response.get('DocumentMetadata', {}).get('Pages', [{}])[0]
        if 'DocumentLocation' in job_response:
            original_key = job_response['DocumentLocation']['S3Object']['Name']
        else:
            # Fallback: assume PDF is in input folder
            original_key = f"input/{job_id}.pdf"
        
        # Define new key in processed_pdfs folder
        filename = original_key.split('/')[-1]  # Get just the filename
        new_key = f"processed_pdfs/{filename}"
        
        # Copy file to new location
        s3.copy_object(
            Bucket=bucket_name,
            CopySource={'Bucket': bucket_name, 'Key': original_key},
            Key=new_key
        )
        
        # Delete original file
        s3.delete_object(Bucket=bucket_name, Key=original_key)
        
        print(f"Moved PDF from {original_key} to {new_key}")
        
    except Exception as e:
        print(f"Failed to move PDF: {str(e)}")
        # Don't fail the entire process if PDF move fails